from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──
section = doc.sections[0]
section.top_margin    = Cm(2.0)
section.bottom_margin = Cm(2.0)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)

# ── Helper: set paragraph shading ──
def shade_paragraph(para, hex_color):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    pPr.append(shd)

def shade_cell(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top','left','bottom','right'):
        tag  = OxmlElement(f'w:{edge}')
        tag.set(qn('w:val'),   kwargs.get('val',  'single'))
        tag.set(qn('w:sz'),    kwargs.get('sz',   '6'))
        tag.set(qn('w:space'), '0')
        tag.set(qn('w:color'), kwargs.get('color','1F4E79'))
        tcBorders.append(tag)
    tcPr.append(tcBorders)

def add_heading(doc, text, level=1, color='1F4E79', size=14, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, shade=None):
    p = doc.add_paragraph()
    p.alignment = align
    if shade:
        shade_paragraph(p, shade)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.font.name = 'Mangal'
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'),    'Mangal')
    rFonts.set(qn('w:hAnsi'),    'Mangal')
    rFonts.set(qn('w:cs'),       'Mangal')
    rPr.insert(0, rFonts)
    return p

def add_para(doc, text, size=11, bold=False, color='000000', indent=False, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.space_before = Pt(2)
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.font.name = 'Mangal'
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'),    'Mangal')
    rFonts.set(qn('w:hAnsi'),    'Mangal')
    rFonts.set(qn('w:cs'),       'Mangal')
    rPr.insert(0, rFonts)
    return p

def add_bullet(doc, text, size=11, color='000000', level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.left_indent  = Cm(0.5 + level * 0.5)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.font.name = 'Mangal'
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'),    'Mangal')
    rFonts.set(qn('w:hAnsi'),    'Mangal')
    rFonts.set(qn('w:cs'),       'Mangal')
    rPr.insert(0, rFonts)
    return p

def add_table_row(table, col1, col2, header=False, shade1='D6E4F0', shade2='EAF4FB'):
    row  = table.add_row()
    c1, c2 = row.cells[0], row.cells[1]
    c1.text = col1
    c2.text = col2
    for c, sh in [(c1, shade1), (c2, shade2)]:
        shade_cell(c, sh)
        set_cell_border(c)
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for para in c.paragraphs:
            para.paragraph_format.space_before = Pt(3)
            para.paragraph_format.space_after  = Pt(3)
            for run in para.runs:
                run.font.name = 'Mangal'
                run.font.size = Pt(10)
                run.bold = header
                run.font.color.rgb = RGBColor.from_string('1F4E79' if header else '000000')
                r = run._r
                rPr = r.get_or_add_rPr()
                rFonts = OxmlElement('w:rFonts')
                rFonts.set(qn('w:ascii'), 'Mangal')
                rFonts.set(qn('w:hAnsi'), 'Mangal')
                rFonts.set(qn('w:cs'),    'Mangal')
                rPr.insert(0, rFonts)
    return row

def divider(doc, color='1F4E79'):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.space_before = Pt(0)


# ════════════════════════════════════════════════
#  TITLE BLOCK
# ════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
shade_paragraph(p, '1F4E79')
run = p.add_run('बीमा लोकपाल (Insurance Ombudsman) — शिकायत प्रस्तुति पत्र')
run.bold = True
run.font.size = Pt(15)
run.font.color.rgb = RGBColor.from_string('FFFFFF')
run.font.name = 'Mangal'
r = run._r
rPr = r.get_or_add_rPr()
rFonts = OxmlElement('w:rFonts')
rFonts.set(qn('w:ascii'), 'Mangal'); rFonts.set(qn('w:hAnsi'), 'Mangal'); rFonts.set(qn('w:cs'), 'Mangal')
rPr.insert(0, rFonts)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
shade_paragraph(p2, '2E75B6')
run2 = p2.add_run('ACKO General Insurance द्वारा Health Claim की गलत अस्वीकृति के विरुद्ध')
run2.bold = True
run2.font.size = Pt(12)
run2.font.color.rgb = RGBColor.from_string('FFFFFF')
run2.font.name = 'Mangal'
r2 = run2._r
rPr2 = r2.get_or_add_rPr()
rFonts2 = OxmlElement('w:rFonts')
rFonts2.set(qn('w:ascii'), 'Mangal'); rFonts2.set(qn('w:hAnsi'), 'Mangal'); rFonts2.set(qn('w:cs'), 'Mangal')
rPr2.insert(0, rFonts2)

doc.add_paragraph().paragraph_format.space_after = Pt(2)


# ════════════════════════════════════════════════
#  SECTION 1 — CASE DETAILS TABLE
# ════════════════════════════════════════════════
add_heading(doc, '1.  मामले का विवरण (Case Details)', size=13, shade='D6E4F0')

table = doc.add_table(rows=0, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = 'Table Grid'
table.columns[0].width = Cm(6)
table.columns[1].width = Cm(10)

rows_data = [
    ('विषय', 'विवरण'),
    ('Policy (पॉलिसी) संख्या', 'ASHP202400001607'),
    ('Claim (दावा) संख्या', '25062100350/1'),
    ('Ombudsman Complaint (लोकपाल शिकायत) संख्या', 'NOI-H-056-2526-0961'),
    ('बीमाधारक का नाम', 'विकास गांधी'),
    ('पता', 'D-20, प्रथम तल, दयानंद नगर, गाजियाबाद, उत्तर प्रदेश – 201001'),
    ('Hospitalization (अस्पताल में भर्ती) दिनांक', '20 जून 2025 (भर्ती) – 21 जून 2025 (छुट्टी)'),
    ('Diagnosis (रोग-निदान)', 'Tight Phimosis (कसा हुआ चमड़ी संकुचन) + Balanitis (शिश्न की सूजन) + Meatal Stenosis (मूत्रमार्ग का सिकुड़ना)'),
    ('Surgery (शल्य-चिकित्सा)', 'Circumcision (खतना) + Meatoplasty (मूत्रमार्ग विस्तार शल्यक्रिया)'),
    ('Biopsy (ऊतक-परीक्षण) निष्कर्ष', 'Lichen Sclerosus / Balanitis Xerotica Obliterans — घातक रोग (Cancer) से मुक्त'),
    ('Claim (दावा) राशि', '₹37,358'),
    ('ACKO द्वारा Rejection (अस्वीकृति) दिनांक', '26 जून 2025 (प्रथम) / 2 जुलाई 2025 (अंतिम)'),
    ('माँगी गई राहत', '₹37,358 (Reimbursement / प्रतिपूर्ति) + ₹10,000 (Mental Agony / मानसिक पीड़ा हेतु क्षतिपूर्ति)'),
]

for i, (c1, c2) in enumerate(rows_data):
    if i == 0:
        add_table_row(table, c1, c2, header=True, shade1='1F4E79', shade2='1F4E79')
        # fix header text color white
        for cell in table.rows[0].cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.color.rgb = RGBColor.from_string('FFFFFF')
    else:
        add_table_row(table, c1, c2)

doc.add_paragraph().paragraph_format.space_after = Pt(4)


# ════════════════════════════════════════════════
#  SECTION 2 — ACKO'S WRONG ARGUMENT
# ════════════════════════════════════════════════
add_heading(doc, '2.  ACKO का गलत तर्क क्या था?', size=13, shade='D6E4F0')

add_para(doc, 'ACKO ने Claim (दावा) अस्वीकार करते हुए केवल यह कहा कि:', size=11)
p_box = doc.add_paragraph()
shade_paragraph(p_box, 'FFF2CC')
p_box.paragraph_format.left_indent  = Cm(0.5)
p_box.paragraph_format.right_indent = Cm(0.5)
p_box.paragraph_format.space_after  = Pt(2)
run_box = p_box.add_run('"आपकी बीमारी Clause (खंड) 4.1.2 — Specific Illnesses (विशिष्ट रोग) की Waiting Period (प्रतीक्षा अवधि) के अंतर्गत आती है, इसलिए 2 वर्षों तक यह खर्च covered (आच्छादित) नहीं होगा।"')
run_box.italic = True
run_box.font.size = Pt(11)
run_box.font.color.rgb = RGBColor.from_string('7B3F00')
run_box.font.name = 'Mangal'
r = run_box._r; rPr = r.get_or_add_rPr(); rFonts = OxmlElement('w:rFonts')
rFonts.set(qn('w:ascii'),'Mangal'); rFonts.set(qn('w:hAnsi'),'Mangal'); rFonts.set(qn('w:cs'),'Mangal'); rPr.insert(0,rFonts)

add_para(doc, 'यह तर्क कानूनी और तथ्यात्मक दोनों दृष्टि से गलत है। नीचे बिंदुवार खंडन प्रस्तुत है:', size=11, bold=True, color='C00000')

doc.add_paragraph().paragraph_format.space_after = Pt(2)


# ════════════════════════════════════════════════
#  SECTION 3 — 5 STRONG ARGUMENTS
# ════════════════════════════════════════════════
add_heading(doc, '3.  हमारे 5 प्रमुख तर्क (Our 5 Strong Arguments)', size=13, shade='D6E4F0')

# ARG 1
add_heading(doc, 'तर्क 1 — बीमारी Clause 4.1.2 की List में है ही नहीं', size=12, color='C00000', shade='FFE7E7')
add_para(doc,'Clause (खंड) 4.1.2 केवल उन्हीं बीमारियों पर लागू होता है जो Policy (पॉलिसी) की Schedule (अनुसूची) में नाम से listed (सूचीबद्ध) हों।', size=11)
add_bullet(doc, 'Tight Phimosis (कसा हुआ चमड़ी संकुचन), Balanitis (शिश्न की सूजन) और Meatal Stenosis (मूत्रमार्ग संकुचन) — ये तीनों ACKO की किसी भी List में नहीं हैं।')
add_bullet(doc, 'Ombudsman (लोकपाल) के समक्ष कहें: "ACKO अपनी Policy Document (पॉलिसी दस्तावेज़) से दिखाए कि इस बीमारी का नाम List में कहाँ है — वह ऐसा नहीं कर पाएगा।"')
add_bullet(doc, 'IRDAI (भारतीय बीमा विनियामक एवं विकास प्राधिकरण) के नियमों के अनुसार, Exclusion (बहिष्करण) स्पष्ट और नाम से लिखित होनी चाहिए।')

divider(doc)

# ARG 2
add_heading(doc, 'तर्क 2 — यह Emergency Surgery (आपातकालीन शल्यक्रिया) थी, Elective (वैकल्पिक) नहीं', size=12, color='C00000', shade='FFE7E7')
add_para(doc,'Medical Emergency (चिकित्सा आपातकाल) की स्थिति में किसी भी Waiting Period (प्रतीक्षा अवधि) का नियम लागू नहीं होता।', size=11)
add_bullet(doc, 'Urinary Obstruction (मूत्र अवरोध) हो रहा था — पेशाब रुक रहा था। यह जानलेवा स्थिति थी।')
add_bullet(doc, 'Doctor (चिकित्सक) ने Immediate Surgery (तत्काल शल्यचिकित्सा) की लिखित सिफारिश की थी।')
add_bullet(doc, 'यह Cosmetic (सौंदर्य-वर्धक) या Optional (ऐच्छिक) Surgery नहीं थी — यह जीवन रक्षक प्रक्रिया थी।')
add_bullet(doc, 'Biopsy Report (ऊतक-परीक्षण रिपोर्ट) ने Lichen Sclerosus की पुष्टि की — जो एक गंभीर त्वचा रोग है।')

divider(doc)

# ARG 3
add_heading(doc, 'तर्क 3 — यह Pre-Existing Disease (पूर्व-विद्यमान रोग) नहीं था', size=12, color='C00000', shade='FFE7E7')
add_para(doc,'ACKO यह सिद्ध नहीं कर सकता कि यह बीमारी Policy (पॉलिसी) लेने से पहले मौजूद थी।', size=11)
add_bullet(doc, 'कोई पुरानी Medical Record (चिकित्सा अभिलेख) नहीं है जिसमें यह बीमारी दर्ज हो।')
add_bullet(doc, 'Proposal Form (प्रस्ताव पत्र) में इसे Disclose (प्रकट) नहीं किया गया — क्योंकि था ही नहीं।')
add_bullet(doc, 'Clause 4.1.2(d) भी केवल Listed Conditions (सूचीबद्ध स्थितियों) पर लागू होता है — और यह बीमारी List में नहीं है।')

divider(doc)

# ARG 4
add_heading(doc, 'तर्क 4 — Policy (पॉलिसी) पूर्ण रूप से Active (सक्रिय) थी', size=12, color='C00000', shade='FFE7E7')
add_para(doc,'Claim (दावे) की तारीख को Policy पूर्णतः वैध और सक्रिय थी।', size=11)
add_bullet(doc, 'जून 2025 तक की सभी Installments (किश्तें / EMI) समय पर जमा की गई थीं।')
add_bullet(doc, 'Policy में कोई Break (अंतराल) नहीं था — यह Fresh Policy (नई पॉलिसी) थी।')
add_bullet(doc, 'Payment Receipts (भुगतान रसीदें) सबूत के तौर पर प्रस्तुत की जा सकती हैं।')

divider(doc)

# ARG 5
add_heading(doc, 'तर्क 5 — ACKO ने Natural Justice (प्राकृतिक न्याय) का उल्लंघन किया', size=12, color='C00000', shade='FFE7E7')
add_para(doc,'ACKO ने बिना उचित जाँच के Claim Reject (दावा अस्वीकार) कर दिया।', size=11)
add_bullet(doc, 'कोई Medical Expert (चिकित्सा विशेषज्ञ) नहीं बुलाया गया।')
add_bullet(doc, 'Hospital (अस्पताल) या Doctor से कोई Verification (सत्यापन) नहीं किया गया।')
add_bullet(doc, 'केवल एक Generic Rejection Letter (सामान्य अस्वीकृति पत्र) भेज दिया — बिना विशिष्ट कारण बताए।')
add_bullet(doc, 'IRDAI के Regulations (विनियम) के अनुसार Insurer (बीमाकर्ता) को Rejection से पहले Claimant (दावेदार) को सुनने का अवसर देना अनिवार्य है।')

doc.add_paragraph().paragraph_format.space_after = Pt(4)


# ════════════════════════════════════════════════
#  SECTION 4 — DOCUMENTS
# ════════════════════════════════════════════════
add_heading(doc, '4.  साथ लाने वाले Documents (दस्तावेज़)', size=13, shade='D6E4F0')
add_para(doc,'Ombudsman (लोकपाल) की Hearing (सुनवाई) में नीचे दिए सभी Documents अवश्य लाएँ:', size=11)

docs = [
    'Hospital Discharge Summary (अस्पताल छुट्टी सारांश) — 20-21 जून 2025',
    'Biopsy Report (ऊतक-परीक्षण रिपोर्ट) — Lichen Sclerosus की पुष्टि',
    'Doctor\'s Prescription (चिकित्सक का पर्चा) — तत्काल Surgery की सिफारिश',
    'Hospital Bill (अस्पताल बिल) — ₹37,358 का पूरा विवरण',
    'Policy Document (पॉलिसी दस्तावेज़) — ASHP202400001607',
    'EMI / Premium Payment Receipts (प्रीमियम भुगतान रसीदें)',
    'ACKO का Rejection Letter (अस्वीकृति पत्र) — 26 जून व 2 जुलाई 2025',
    'सभी Emails (ईमेल) जो ACKO को भेजी गईं और उनके जवाब',
    'Ombudsman Complaint (लोकपाल शिकायत) की प्रति — NOI-H-056-2526-0961',
    'Annexure VI-A Form (अनुलग्नक VI-A प्रपत्र) — भरा हुआ',
]
for d in docs:
    add_bullet(doc, d, color='1F4E79')

doc.add_paragraph().paragraph_format.space_after = Pt(4)


# ════════════════════════════════════════════════
#  SECTION 5 — HEARING MEIN KYA KAHEIN
# ════════════════════════════════════════════════
add_heading(doc, '5.  Hearing (सुनवाई) में क्या कहें — तैयार वाक्य', size=13, shade='D6E4F0')
add_para(doc,'नीचे दिए वाक्य Ombudsman के सामने आत्मविश्वास से बोलें:', size=11, bold=True)

statements = [
    '"Clause (खंड) 4.1.2 की List में Phimosis / Balanitis का नाम नहीं है — ACKO अपनी Policy Document से दिखाए।"',
    '"यह Emergency (आपातकाल) था — Doctor ने Immediate Surgery (तत्काल शल्यचिकित्सा) लिखी थी — Elective (वैकल्पिक) नहीं था।"',
    '"Biopsy Report (ऊतक-परीक्षण रिपोर्ट) से Lichen Sclerosus सिद्ध है — यह Pre-Existing Disease (पूर्व-विद्यमान रोग) नहीं थी।"',
    '"मेरी सभी EMIs (किश्तें) समय पर जमा हैं — Policy पूरी तरह Active (सक्रिय) थी।"',
    '"ACKO ने कोई Medical Expert (विशेषज्ञ) नहीं बुलाया — यह IRDAI Guidelines (दिशानिर्देश) का उल्लंघन है।"',
    '"मैं ₹37,358 की Reimbursement (प्रतिपूर्ति) और ₹10,000 Mental Agony Compensation (मानसिक पीड़ा क्षतिपूर्ति) का हकदार हूँ।"',
]
for s in statements:
    p = add_bullet(doc, s, color='1A5276')
    for run in p.runs:
        run.bold = True

doc.add_paragraph().paragraph_format.space_after = Pt(4)


# ════════════════════════════════════════════════
#  SECTION 6 — ACKO KA POSSIBLE COUNTER
# ════════════════════════════════════════════════
add_heading(doc, '6.  ACKO का संभावित Counter-Argument (जवाबी तर्क) और हमारा उत्तर', size=13, shade='D6E4F0')

table2 = doc.add_table(rows=0, cols=2)
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
table2.style = 'Table Grid'
table2.columns[0].width = Cm(7)
table2.columns[1].width = Cm(9)

counter_rows = [
    ('ACKO का तर्क', 'हमारा उत्तर'),
    ('"यह Specific Illness (विशिष्ट रोग) है — Waiting Period लागू है"', '"List दिखाओ जिसमें Phimosis/Balanitis का नाम हो — है ही नहीं।"'),
    ('"Policy नई है, 2 साल पूरे नहीं हुए"', '"Waiting Period केवल Listed Diseases पर लागू होती है — यह Listed नहीं है।"'),
    ('"2022 की Ultrasound (अल्ट्रासाउंड) रिपोर्ट में Prostate (प्रोस्टेट) बड़ा था"', '"Prostate की बीमारी और Phimosis/Balanitis दो अलग-अलग Organs की अलग-अलग बीमारियाँ हैं — कोई Medical Connection (चिकित्सीय संबंध) नहीं।"'),
    ('"हमने Senior Committee (वरिष्ठ समिति) से Review (समीक्षा) करवाई"', '"Committee का नाम और Medical Expert का नाम बताएँ — कोई Hospital Verification नहीं हुई।"'),
]

for i, (c1, c2) in enumerate(counter_rows):
    if i == 0:
        add_table_row(table2, c1, c2, header=True, shade1='1F4E79', shade2='1F4E79')
        for cell in table2.rows[0].cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.color.rgb = RGBColor.from_string('FFFFFF')
    else:
        add_table_row(table2, c1, c2, shade1='FFE7E7', shade2='E8F8E8')

doc.add_paragraph().paragraph_format.space_after = Pt(4)


# ════════════════════════════════════════════════
#  SECTION 7 — CONCLUSION
# ════════════════════════════════════════════════
add_heading(doc, '7.  निष्कर्ष (Conclusion)', size=13, shade='D6E4F0')

p_conc = doc.add_paragraph()
shade_paragraph(p_conc, 'E8F5E9')
p_conc.paragraph_format.left_indent  = Cm(0.3)
p_conc.paragraph_format.right_indent = Cm(0.3)
p_conc.paragraph_format.space_after  = Pt(4)
run_conc = p_conc.add_run(
    'आपका Case (मामला) कानूनी और तथ्यात्मक दोनों दृष्टि से मज़बूत है। '
    'ACKO ने गलत Clause (खंड) लगाकर एक जरूरी और आपातकालीन Surgery (शल्यचिकित्सा) के Claim (दावे) को ठुकरा दिया। '
    'आपके पास Hospital Discharge Summary, Biopsy Report, Doctor Prescriptions, और EMI Payment Records — '
    'सभी प्रमाण मौजूद हैं। '
    'Ombudsman (लोकपाल) के सामने शांति और आत्मविश्वास से अपनी बात रखें — '
    'जीतने की पूरी संभावना है।'
)
run_conc.font.size = Pt(11)
run_conc.font.color.rgb = RGBColor.from_string('1A5276')
run_conc.font.name = 'Mangal'
r = run_conc._r; rPr = r.get_or_add_rPr(); rFonts = OxmlElement('w:rFonts')
rFonts.set(qn('w:ascii'),'Mangal'); rFonts.set(qn('w:hAnsi'),'Mangal'); rFonts.set(qn('w:cs'),'Mangal')
rPr.insert(0, rFonts)


# ════════════════════════════════════════════════
#  FOOTER
# ════════════════════════════════════════════════
footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
shade_paragraph(footer_p, '1F4E79')
footer_run = footer_p.add_run('विकास गांधी  |  Policy: ASHP202400001607  |  Complaint: NOI-H-056-2526-0961  |  Mobile: 9911198000')
footer_run.font.size = Pt(9)
footer_run.font.color.rgb = RGBColor.from_string('FFFFFF')
footer_run.font.name = 'Mangal'
r = footer_run._r; rPr = r.get_or_add_rPr(); rFonts = OxmlElement('w:rFonts')
rFonts.set(qn('w:ascii'),'Mangal'); rFonts.set(qn('w:hAnsi'),'Mangal'); rFonts.set(qn('w:cs'),'Mangal')
rPr.insert(0, rFonts)


# ── Save ──
out = '/home/user/vikas/ACKO_STRUCTURED/Documents/Word Processing/Ombudsman_Case_HindiGuide.docx'
doc.save(out)
print('SAVED:', out)
